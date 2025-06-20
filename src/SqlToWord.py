import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm


def clean_quoted_identifier(identifier):
    """清除标识符的引号并处理转义字符"""
    if len(identifier) < 2:
        return identifier

    # 检查是否被引号包裹
    if (identifier[0] == '`' and identifier[-1] == '`') or \
            (identifier[0] == '"' and identifier[-1] == '"') or \
            (identifier[0] == "'" and identifier[-1] == "'"):
        # 获取引号类型
        quote_char = identifier[0]

        # 提取内容并处理转义引号
        inner = identifier[1:-1]
        inner = inner.replace(quote_char * 2, quote_char)  # 替换双引号为单引号
        return inner

    # 无引号标识符直接返回
    return identifier


def extract_table_comment(table_options):
    """从表选项中提取表备注信息"""
    if not table_options:
        return ''

    # 匹配 COMMENT 选项（支持单引号和双引号）
    comment_pattern = r'COMMENT\s*=\s*(([\'"])(.*?)(?<!\\)\2)'
    match = re.search(comment_pattern, table_options, re.IGNORECASE | re.DOTALL)

    if match:
        comment = match.group(3)
        # 处理转义引号：将两个连续引号替换为单个引号
        quote_char = match.group(2)
        if quote_char:
            comment = comment.replace(quote_char * 2, quote_char)
        return comment

    # 尝试匹配无引号的 COMMENT（不常见但可能）
    unquoted_comment_pattern = r'COMMENT\s*=\s*([^\s,]+)'
    match = re.search(unquoted_comment_pattern, table_options, re.IGNORECASE)
    if match:
        return match.group(1)

    # 尝试匹配旧式注释（不带等号）
    old_style_pattern = r'COMMENT\s*([\'"])(.*?)(?<!\\)\1'
    match = re.search(old_style_pattern, table_options, re.IGNORECASE | re.DOTALL)
    if match:
        comment = match.group(2)
        quote_char = match.group(1)
        if quote_char:
            comment = comment.replace(quote_char * 2, quote_char)
        return comment

    return ''


def parse_fields(create_definition):
    """
    从 CREATE TABLE 语句的列定义部分解析字段信息
    支持无换行符的连续字段定义
    :param create_definition: CREATE TABLE 语句中的列定义部分
    :return: 字段对象列表
    """
    if not create_definition:
        return []

    # 改进的字段定义正则表达式，处理连续字段定义
    field_pattern = re.compile(
        r"`?(?P<name>[^`\s,]+)`?"  # 字段名（可能被反引号包裹）
        r"\s+"
        r"(?P<type>\w+)"  # 字段类型
        r"(?:\s*\(\s*(?P<length>[^)]+)\s*\))?"  # 长度/精度（可选）
        r"(?:\s+(?P<nullable>NOT\s+NULL|NULL))?"  # 是否可为空（可选）
        r".*?"  # 其他字段属性
        r"(?:COMMENT\s+'(?P<comment>(?:''|[^'])*)')?"  # 字段注释（可选）
        r"(?=,|$)",  # 前瞻断言，确保匹配到下一个字段或结束
        re.IGNORECASE | re.DOTALL
    )

    fields = []
    pos = 0

    # 跳过主键、索引等定义行
    skip_pattern = re.compile(
        r"^\s*(PRIMARY\s+KEY|KEY|CONSTRAINT|UNIQUE|INDEX|FOREIGN\s+KEY|CHECK)\b",
        re.IGNORECASE
    )

    # 处理连续字段定义
    while pos < len(create_definition):
        # 检查是否跳过非字段行
        if skip_pattern.match(create_definition[pos:]):
            # 找到下一个逗号或结束位置
            next_comma = create_definition.find(",", pos)
            if next_comma == -1:
                break
            pos = next_comma + 1
            continue

        # 尝试匹配字段定义
        field_match = field_pattern.search(create_definition, pos)
        if not field_match:
            # 如果没有匹配，移动到下一个逗号
            next_comma = create_definition.find(",", pos)
            if next_comma == -1:
                break
            pos = next_comma + 1
            continue

        # 提取字段信息
        name = field_match.group('name')
        if name.startswith('`') and name.endswith('`'):
            name = name[1:-1]

        data_type = field_match.group('type').upper()

        if field_match.group('length'):
            data_type += f"({field_match.group('length').strip()})"

        nullable = 'YES'  # 默认可为空
        if field_match.group('nullable'):
            nullable = 'NO' if 'NOT NULL' in field_match.group('nullable').upper() else 'YES'

        comment = field_match.group('comment') or ''
        if comment:
            comment = comment.replace("''", "'")

        fields.append({
            'name': name,
            'type': data_type,
            'nullable': nullable,
            'comment': comment
        })

        # 移动到匹配结束位置
        pos = field_match.end()

        # 跳过逗号（如果有）
        if pos < len(create_definition) and create_definition[pos] == ',':
            pos += 1

    return fields


def parse_create_table(create_content):
    """
    解析 SQL 内容，提取表结构信息
    :param create_content: CREATE TABLE 语句的内容
    :return: 包含表结构信息的对象
    """

    # 1. 提取 CREATE [TEMPORARY] TABLE [IF NOT EXISTS]
    header_pattern = r'CREATE\s+(?P<temporary>TEMPORARY\s+)?TABLE\s+(?P<if_not_exists>IF\s+NOT\s+EXISTS\s+)?'
    header_match = re.match(header_pattern, create_content, re.IGNORECASE)
    if not header_match:
        return None

    # 提取头部组件
    temporary = header_match.group('temporary') or ''
    if_not_exists = header_match.group('if_not_exists') or ''

    # 2. 提取表名 (支持四种格式)
    remaining = create_content[header_match.end():].lstrip()
    table_name_pattern = r'''(?:
            `(?:[^`]|``)+`    |  # 反引号包裹
            '(?:[^']|'')+'    |  # 单引号包裹
            [\w\.]+             # 无引号标识符
        )\s*'''

    table_name_match = re.search(table_name_pattern, remaining, re.VERBOSE)
    if not table_name_match:
        return None

    table_name = clean_quoted_identifier(table_name_match.group().strip())
    remaining = remaining[table_name_match.end():].lstrip()

    # 3. 提取列定义 [(create_definition,...)]
    create_definition = ''
    if remaining.startswith('('):
        # 使用括号匹配处理嵌套结构
        level = 0
        in_string = False
        escape_next = False
        end_index = 0

        for i, char in enumerate(remaining):
            if escape_next:
                escape_next = False
                continue

            if char == '\\':
                escape_next = True
            elif char in ('"', "'"):
                in_string = not in_string
            elif not in_string:
                if char == '(':
                    level += 1
                elif char == ')':
                    level -= 1
                    if level == 0:
                        end_index = i
                        break

        if level != 0:
            return None  # 括号不匹配

        create_definition = remaining[1:end_index].strip().replace('\n', '').replace('\r', '')
        remaining = remaining[end_index + 1:].strip()

    # 4. 提取表选项 [table_options]
    # 表选项通常以分号或 PARTITION 关键字结束
    table_options = ''
    if remaining and not re.match(r'(?:PARTITION|IGNORE|REPLACE|AS)\b', remaining, re.IGNORECASE):
        options_match = re.search(
            r'^(.*?)(?=(?:\s+PARTITION\s+|\s+IGNORE\s+|\s+REPLACE\s+|\s+AS\s+|\s*;|$))',
            remaining,
            re.DOTALL | re.IGNORECASE
        )
        if options_match:
            table_options = options_match.group(1).strip()
            remaining = remaining[options_match.end():].strip()

    # 5. 提取分区选项 [partition_options]
    partition_options = ''
    if re.match(r'PARTITION\s+', remaining, re.IGNORECASE):
        partition_match = re.search(
            r'^(PARTITION\s+BY.*?)(?=(?:\s+IGNORE\s+|\s+REPLACE\s+|\s+AS\s+|\s*;|$))',
            remaining,
            re.DOTALL | re.IGNORECASE
        )
        if partition_match:
            partition_options = partition_match.group(1).strip()
            remaining = remaining[partition_match.end():].strip()

    # 6. 提取 IGNORE/REPLACE 和 AS 子句
    ignore_replace = ''
    as_clause = ''
    # 处理 IGNORE | REPLACE
    ignore_match = re.match(r'(IGNORE|REPLACE)\s+', remaining, re.IGNORECASE)
    if ignore_match:
        ignore_replace = ignore_match.group(1).upper()
        remaining = remaining[ignore_match.end():].strip()
    # 处理 AS 子句
    if re.match(r'AS\s+', remaining, re.IGNORECASE):
        as_clause = remaining[2:].strip()

    # 7. 提取表备注信息
    table_comment = extract_table_comment(table_options)

    return {
        'temporary': temporary.strip(),
        'if_not_exists': if_not_exists.strip(),
        'table_name': table_name,
        'table_comment': table_comment,
        'create_definition': create_definition,
        'table_options': table_options,
        'partition_options': partition_options,
        'ignore_replace': ignore_replace,
        'as_clause': as_clause,
        'fields': parse_fields(create_definition)
    }


def parse_sql(sql_content):
    """
    解析 SQL 内容，提取表列表信息
    :param sql_content: SQL 文件的内容
    :return: 包含表结构信息的列表
    """

    # 切分 SQL 获得 CREATE TABLE 语句
    table_pattern = re.compile(
        r'CREATE\s+(?:TEMPORARY\s+)?TABLE\s+(?:IF\s+NOT\s+EXISTS\s+)?'
        r'(?:'
        r'`(?:[^`]|``)+`'  # 匹配反引号表名 (允许内部转义反引号)
        r"|'(?:[^']|'')+'"  # 匹配单引号表名 (允许内部转义单引号)
        r'|\w+(?:\.\w+)*'  # 匹配无引号表名 (允许 db.table 格式)
        r')'
        r'\s*.*?;',  # 匹配表定义直到分号结束
        re.DOTALL | re.IGNORECASE
    )

    tables = []
    # 获取全部 CREATE TABLE 语句 - 存储为 list
    for table in table_pattern.findall(sql_content):
        tables.append(parse_create_table(table))
    return tables


def set_paragraph_style(paragraph, style_info, default_info):
    """
    设置段落样式，支持样式部分不同或为空
    :param paragraph: 段落对象
    :param style_info: 样式信息
    :param default_info: 默认样式信息
    """
    # 获取对齐方式，若样式信息中没有则使用默认值
    align = style_info.get('align', default_info['align'])
    # 获取行间距，若样式信息中没有则使用默认值
    line_spacing = style_info.get('line_spacing', default_info['line_spacing'])
    # 获取段前间距，若样式信息中没有则使用默认值
    space_before = style_info.get('space_before', default_info['space_before'])
    # 获取段后间距，若样式信息中没有则使用默认值
    space_after = style_info.get('space_after', default_info['space_after'])
    # 获取字体，若样式信息中没有则使用默认值
    font = style_info.get('font', default_info['font'])
    # 获取字号，若样式信息中没有则使用默认值
    size = style_info.get('size', default_info['size'])
    # 获取字体颜色，若样式信息中没有则使用默认值
    color = style_info.get('color', default_info['color'])
    # 获取中文字体，若样式信息中没有则使用默认值
    font_cn = style_info.get('font-CN', default_info['font-CN'])
    # 获取首行缩进，若样式信息中没有则使用默认值
    first_line_indent = style_info.get('first_line_indent', default_info['first_line_indent'])

    # 设置段落对齐方式
    paragraph.alignment = align
    # 设置段前间距
    paragraph.paragraph_format.space_before = space_before
    # 设置段后间距
    paragraph.paragraph_format.space_after = space_after
    # 设置首行缩进
    paragraph.paragraph_format.first_line_indent = Pt(size * first_line_indent)

    # 遍历段落中的每个运行对象，设置字体、字号、颜色和中文字体
    for run in paragraph.runs:
        run.font.name = font
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)

    # 判断行间距是固定值还是倍数
    if isinstance(line_spacing, Pt):
        # 固定值行间距
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = line_spacing
    else:
        # 倍数行间距
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph.paragraph_format.line_spacing = line_spacing


def set_cell_border(cell, **kwargs):
    """
    设置单元格边框
    :param cell: 单元格对象
    :param kwargs: 边框设置参数
    """
    # 从单元格对象中获取底层的 XML 元素
    tc = cell._tc
    # 获取单元格的属性元素，如果不存在则创建一个新的
    tcPr = tc.get_or_add_tcPr()
    # 获取单元格的边框元素，如果不存在则创建一个新的
    tcBorders = tcPr.first_child_found_in("w:tcBorders") or OxmlElement('w:tcBorders')
    tcPr.append(tcBorders)

    # 定义需要设置边框的四个方向
    for edge, edge_data in kwargs.items():
        if edge in ('start', 'end', 'top', 'bottom'):
            # 构建当前边框方向的 XML 标签名
            tag = qn(f'w:{edge}')
            # 在 tcBorders 中查找当前边框元素，如果不存在则创建一个新的
            element = tcBorders.find(tag) or OxmlElement(f'w:{edge}')
            # 遍历当前边框方向的设置参数，将其应用到边框元素上
            for key, value in edge_data.items():
                element.set(qn(f'w:{key}'), str(value))
            # 如果当前边框元素不在 tcBorders 中，则将其添加进去
            if element not in tcBorders:
                tcBorders.append(element)


def generate_word_doc(tables, output_path, default_style, head_style, content_style, title_style, table_style,
                      basic_num):
    """
    根据表结构信息生成 docx 文档
    :param tables: 包含表结构信息的列表
    :param output_path: 输出文档的路径
    :param default_style: 默认样式信息
    :param head_style: 标题样式信息
    :param content_style: 内容样式信息
    :param title_style: 表格标题样式信息
    :param table_style: 表格样式信息
    :param basic_num: 基本编号信息
    """
    # 创建文档对象
    doc = Document()

    # 设置纸张大小
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    # 遍历表结构信息
    for index, table_data in enumerate(tables):

        length = len(table_data)

        name = table_data['table_name']

        comment = table_data['table_comment'] if table_data['table_comment'] else ''
        attribute_num = basic_num['attribute_num']
        comment_content = ""
        if table_data['table_comment']:
            text = ""
            if attribute_num < length:
                text = "等"
            for row_idx, field in enumerate(table_data['fields'][:attribute_num]):
                comment_content += field['comment']
                if row_idx != attribute_num - 1:
                    comment_content += "、"
            comment_content += text + f"共{length}个属性"

        chapter_num = basic_num['chapter_num']
        start_table_num = basic_num['start_table_num']

        # 添加标题
        if head_style['add']:
            heading = doc.add_heading(f"{index + 1} {comment}表", head_style['style']['level'])
            set_paragraph_style(heading, head_style.get('style', {}), default_style)

        # 添加内容
        if content_style['add']:
            content = doc.add_paragraph(
                f"{comment}表用于存储{comment}信息，包含{comment_content}。{comment}表如表{chapter_num}-{index + start_table_num}所示。")
            set_paragraph_style(content, content_style.get('style', {}), default_style)

        # 添加表格标题
        if title_style['add']:
            title = doc.add_paragraph(f"表{chapter_num}-{index + start_table_num} {comment}表({name})")
            set_paragraph_style(title, title_style.get('style', {}), default_style)

        # 添加表格
        table = doc.add_table(rows=1, cols=4)
        # 自动调整列宽
        table.autofit = False

        table_type = table_style['style']['table_type']
        border_color = table_style.get('style', {}).get('border_color', RGBColor(0, 0, 0))
        hex_border_color = "{:02X}{:02X}{:02X}".format(border_color[0], border_color[1], border_color[2])
        outer_border = table_style['style']['outer_border']
        inner_border = table_style['style']['inner_border']

        # 获取表格第一行的单元格
        hdr_cells = table.rows[0].cells
        # 定义表头
        headers = ['字段名', '类型', '允许空', '说明']

        # 设置表头
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            set_paragraph_style(hdr_cells[i].paragraphs[0], table_style.get('style', {}), default_style)

            for run in hdr_cells[i].paragraphs[0].runs:
                run.bold = True  # 加粗

            # 设置表头边框
            if table_type == 1:  # 三线表
                borders = {
                    'top': {'val': 'single', 'sz': outer_border, 'color': hex_border_color},
                    'bottom': {'val': 'single', 'sz': inner_border, 'color': hex_border_color},
                }
            else:  # 全线表
                borders = {
                    'top': {'val': 'single', 'sz': outer_border, 'color': hex_border_color},
                    'bottom': {'val': 'single', 'sz': inner_border, 'color': hex_border_color},
                    'start': {'val': 'single', 'sz': outer_border if i == 0 else inner_border,
                              'color': hex_border_color},
                    'end': {'val': 'single', 'sz': outer_border if i == 3 else inner_border, 'color': hex_border_color}
                }
            set_cell_border(hdr_cells[i], **borders)

        # 填充数据
        for row_idx, field in enumerate(table_data['fields']):
            # 添加表格
            row = table.add_row().cells
            # 填充字段名
            row[0].text = field['name']
            # 填充字段类型
            row[1].text = field['type']
            # 填充字段是否允许为空
            row[2].text = field['nullable']
            # 填充字段注释
            row[3].text = field['comment']

            # 设置单元格样式
            for cell in row:
                # 设置单元格段落的对齐方式
                set_paragraph_style(cell.paragraphs[0], table_style.get('style', {}), default_style)

                # 设置单元格边框
                if table_type == 1:  # 三线表
                    if row_idx + 1 == len(table_data['fields']):
                        borders = {
                            'bottom': {'val': 'single', 'sz': outer_border, 'color': hex_border_color},
                        }
                    else:
                        borders = {
                            'top': {'val': 'none', 'sz': 0, 'color': hex_border_color},
                            'bottom': {'val': 'none', 'sz': 0, 'color': hex_border_color},
                            'start': {'val': 'none', 'sz': 0, 'color': hex_border_color},
                            'end': {'val': 'none', 'sz': 0, 'color': hex_border_color}
                        }
                else:  # 全线表
                    borders = {
                        'top': {'val': 'single', 'sz': inner_border, 'color': hex_border_color},
                        'bottom': {'val': 'single',
                                   'sz': outer_border if row_idx + 1 == len(table_data['fields']) else inner_border,
                                   'color': hex_border_color},
                        'start': {'val': 'single', 'sz': outer_border if cell == row[0] else inner_border,
                                  'color': hex_border_color},
                        'end': {'val': 'single', 'sz': outer_border if cell == row[-1] else inner_border,
                                'color': hex_border_color}
                    }
                set_cell_border(cell, **borders)

        # 设置列宽
        widths = [Cm(4), Cm(3.2), Cm(2.4), Cm(5)]
        for i in range(len(widths)):
            table.columns[i].width = widths[i]
            # 为整列设置垂直居中（确保所有单元格都居中）
            for cell in table.columns[i].cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), "center")  # 垂直居中
                tcPr.append(vAlign)

        # 表格间是否空行
        if basic_num['blank']:
            doc.add_paragraph()

    # 保存文档
    doc.save(output_path)


def sql_to_word(sql_path, output_path, default_style, head_style, content_style, title_style, table_style, basic_num):
    """
    读取 SQL 文件信息，生成 Word 文档
    :param sql_path: SQL 文件的路径
    :param output_path: 输出 Word 文档的路径
    :param default_style: 默认样式信息
    :param head_style: 标题样式信息
    :param content_style: 内容样式信息
    :param title_style: 表格标题样式信息
    :param table_style: 表格样式信息
    :param basic_num: 基本编号信息
    """
    # 检查 SQL 文件是否存在
    if not os.path.exists(sql_path):
        raise FileNotFoundError(f"SQL文件不存在：{sql_path}")
    # 读取 SQL 文件内容
    with open(sql_path, 'r', encoding='utf-8') as f:
        sql_content = f.read()
    # 生成 Word 文档
    generate_word_doc(parse_sql(sql_content), output_path, default_style, head_style, content_style, title_style,
                      table_style, basic_num)


if __name__ == "__main__":
    sql_to_word(
        sql_path='input.sql',
        output_path='output.docx',
        default_style={
            'font': 'Times New Roman',
            'font-CN': '宋体',
            'size': 12,  # 小四
            'color': RGBColor(0, 0, 0),
            'align': WD_ALIGN_PARAGRAPH.LEFT,  # 左对齐
            'line_spacing': Pt(22),  # 固定值22磅
            'space_before': 0,
            'space_after': 0,
            'first_line_indent': 0,
        },
        head_style={
            'add': True,
            'style': {
                'level': 1,
                'font-CN': '黑体',
                'size': 15,  # 小三
                'line_spacing': 1.5,  # 1.5倍行距
            }
        },
        content_style={
            'add': True,
            'style': {
                'first_line_indent': 2,  # 首行缩进2字符
            }
        },
        title_style={
            'add': True,
            'style': {
                'font-CN': '黑体',
                'size': 10.5,  # 五号
                'align': WD_ALIGN_PARAGRAPH.CENTER,  # 居中对齐
            }
        },
        table_style={
            'add': True,
            'style': {
                'size': 10.5,
                'line_spacing': 1,
                'outer_border': 15,  # 外框线宽度
                'inner_border': 7.5,  # 内部线宽度
                'border_color': RGBColor(0, 0, 0),
                'align': WD_ALIGN_PARAGRAPH.CENTER,
                'table_type': 1,  # 1-三线表 2-全线表
            }
        },
        basic_num={
            'blank': False,  # 是否空行
            'chapter_num': 1,  # 章节编号
            'start_table_num': 1,  # 表格开始标号
            'attribute_num': 3  # 文本属性数量
        }
    )
